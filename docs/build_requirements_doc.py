from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "BASTREET_Requisitos_e_Banco_de_Dados.docx"
DIAGRAM = ROOT / "database" / "diagrama-er.png"
ORANGE = "F37021"; BLACK = "141414"; DARK = "252525"; LIGHT = "F3F4F6"; GRAY = "616161"; WHITE = "FFFFFF"

def font(size, bold=False):
    try: return ImageFont.truetype("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf", size)
    except: return ImageFont.load_default()

def diagram():
    groups = [
        ("IDENTIDADE", ["users", "player_profiles", "player_availability"]),
        ("LOCAIS E TREINOS", ["courts", "training_catalog", "training_sessions", "court_checkins"]),
        ("PARTIDAS", ["events", "event_participants", "matchmaking_queue", "matches", "match_players"]),
        ("RANKING E SOCIAL", ["academic_semesters", "ranking_entries", "conversations", "conversation_members", "messages"]),
    ]
    im = Image.new("RGB", (1800, 1120), "#F7F7F7"); d = ImageDraw.Draw(im)
    d.rounded_rectangle((35, 35, 1765, 1085), 28, fill="#101010")
    d.text((85, 65), "BASTREET — DIAGRAMA ENTIDADE-RELACIONAMENTO", font=font(38, True), fill="#FFFFFF")
    d.text((85, 115), "Modelo PostgreSQL de referência • relacionamentos principais", font=font(24), fill="#C9C9C9")
    positions = {}; xcols = [85, 520, 955, 1390]
    for col, (title, items) in enumerate(groups):
        x=xcols[col]; d.text((x, 185), title, font=font(22, True), fill="#F37021")
        for row, name in enumerate(items):
            y=235+row*145; positions[name]=(x,y)
            d.rounded_rectangle((x,y,x+320,y+82),14,fill="#272727",outline="#F37021",width=3)
            d.text((x+18,y+24),name,font=font(21,True),fill="#FFFFFF")
    edges = [
      ("users","player_profiles"),("users","player_availability"),("users","training_sessions"),
      ("training_catalog","training_sessions"),("courts","training_sessions"),("courts","court_checkins"),
      ("users","events"),("courts","events"),("events","event_participants"),("users","event_participants"),
      ("users","matchmaking_queue"),("events","matches"),("matches","match_players"),("users","match_players"),
      ("academic_semesters","ranking_entries"),("users","ranking_entries"),("events","conversations"),
      ("conversations","conversation_members"),("users","conversation_members"),("conversations","messages")]
    # Conectores discretos; o Mermaid versionado contém a topologia completa e navegável.
    for a,b in edges:
        if a not in positions or b not in positions: continue
        ax,ay=positions[a]; bx,by=positions[b]
        p1=(ax+320,ay+41); p2=(bx,by+41)
        if ax==bx: p1=(ax+160,ay+82); p2=(bx+160,by)
        d.line((p1,p2),fill="#797979",width=2)
        d.ellipse((p2[0]-4,p2[1]-4,p2[0]+4,p2[1]+4),fill="#F37021")
    d.text((85,1033),"Detalhes de cardinalidade e restrições: docs/database/diagrama-er.md e schema.sql",font=font(18),fill="#BDBDBD")
    im.save(DIAGRAM, quality=95)

def shade(cell, color):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn('w:shd'))
    if shd is None: shd=OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'),color)

def set_cell_width(cell, dxa):
    tcPr=cell._tc.get_or_add_tcPr(); tcW=tcPr.find(qn('w:tcW'))
    if tcW is None: tcW=OxmlElement('w:tcW'); tcPr.append(tcW)
    tcW.set(qn('w:w'),str(dxa)); tcW.set(qn('w:type'),'dxa')

def hyperlink(p, text, url):
    part=p.part; rid=part.relate_to(url,'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',is_external=True)
    h=OxmlElement('w:hyperlink'); h.set(qn('r:id'),rid); r=OxmlElement('w:r'); rPr=OxmlElement('w:rPr')
    color=OxmlElement('w:color'); color.set(qn('w:val'),ORANGE); rPr.append(color)
    under=OxmlElement('w:u'); under.set(qn('w:val'),'single'); rPr.append(under); r.append(rPr)
    t=OxmlElement('w:t'); t.text=text; r.append(t); h.append(r); p._p.append(h)

def add_page_number(p):
    p.alignment=WD_ALIGN_PARAGRAPH.RIGHT; r=p.add_run('Página ')
    fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); p._p.append(fld)

def add_table(doc, headers, rows, widths):
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    tblPr=t._tbl.tblPr; tblW=tblPr.find(qn('w:tblW')); tblW.set(qn('w:w'),'9360'); tblW.set(qn('w:type'),'dxa')
    ind=OxmlElement('w:tblInd'); ind.set(qn('w:w'),'120'); ind.set(qn('w:type'),'dxa'); tblPr.append(ind)
    grid=t._tbl.tblGrid
    for old in list(grid): grid.remove(old)
    for w in widths: el=OxmlElement('w:gridCol'); el.set(qn('w:w'),str(w)); grid.append(el)
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; set_cell_width(c,widths[i]); shade(c,DARK); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0); r=p.add_run(h); r.bold=True; r.font.color.rgb=RGBColor(255,255,255)
    header_pr=t.rows[0]._tr.get_or_add_trPr(); repeat=OxmlElement('w:tblHeader'); repeat.set(qn('w:val'),'true'); header_pr.append(repeat)
    for row in rows:
        cells=t.add_row().cells
        row_pr=t.rows[-1]._tr.get_or_add_trPr(); no_split=OxmlElement('w:cantSplit'); row_pr.append(no_split)
        for i,val in enumerate(row):
            set_cell_width(cells[i],widths[i]); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.TOP
            if len(t.rows)%2==1: shade(cells[i],LIGHT)
            p=cells[i].paragraphs[0]; p.paragraph_format.space_after=Pt(2); p.add_run(str(val))
    doc.add_paragraph().paragraph_format.space_after=Pt(0)
    return t

RF = [
('RF-01','Cadastro de usuário','Permitir o cadastro com nome, e-mail, senha, idade/data de nascimento, altura, posição, gênero, localização, nível declarado e dias/horários de disponibilidade.'),
('RF-02','Autenticação','Permitir login por e-mail e senha, manter sessão autenticada e encerrar a sessão com segurança.'),
('RF-03','Perfil do jogador','Exibir e permitir atualizar foto, nome, posição, altura, nível técnico, XP, partidas, vitórias, derrotas e pontuação semestral.'),
('RF-04','Localização do dispositivo','Solicitar consentimento para usar o GPS, tratar recusa/indisponibilidade e usar Recife/PE como referência de contingência.'),
('RF-05','Busca de quadras','Buscar quadras próximas à coordenada atual, combinar uma base curada com dados do OpenStreetMap e ordenar os resultados por distância.'),
('RF-06','Mapa e rota','Centralizar o mapa no usuário, mostrar marcadores e detalhes das quadras e oferecer abertura de rota externa até a quadra selecionada.'),
('RF-07','Presença em quadra','Permitir check-in do usuário em uma quadra e apresentar a presença recente de jogadores no local.'),
('RF-08','Modos de partida','Disponibilizar partida normal, rápida e personalizada, com explicação clara do efeito de cada modalidade.'),
('RF-09','Fila de matchmaking','Inserir e remover o usuário da fila, impedir duplicidade e informar quantidade de participantes e estado da busca.'),
('RF-10','Formação de equipes','Cruzar nível técnico, altura, gênero e disponibilidade para sugerir equipes equilibradas, sem usar o ranking semestral como critério técnico.'),
('RF-11','Partida personalizada','Permitir definir quadra, data, horário, regras, capacidade e convites para um evento organizado pelo usuário.'),
('RF-12','Entrada em evento','Permitir entrar em um treino ou partida agendada e atualizar a lista e a quantidade de participantes.'),
('RF-13','Confirmação e resultado','Permitir confirmar presença, registrar o time vencedor e consolidar o resultado da partida.'),
('RF-14','Treinos individuais','Exibir treinos que possam ser feitos em casa ou em qualquer local e registrar sua conclusão diária.'),
('RF-15','Treinos coletivos','Registrar treino em grupo vinculado a uma quadra e conceder pontuação superior à do treino individual.'),
('RF-16','Progressão','Atualizar XP, nível técnico e pontos semestrais após atividades válidas e apresentar retorno imediato ao usuário.'),
('RF-17','Ranking semestral','Exibir classificação por período letivo com filtros de cidade, geral e amigos, preparando a seleção para os campeonatos finais.'),
('RF-18','Chat','Permitir envio e consulta de mensagens entre participantes, preservando autoria e ordem cronológica.'),
('RF-19','Respostas automatizadas','Gerar respostas variadas de bots conforme palavras-chave e contexto básico da mensagem para sustentar a demonstração sem fluxo real de usuários.'),
('RF-20','Administração','Permitir que administradores acompanhem usuários, quadras, eventos e registros relevantes para a demonstração acadêmica.')]

RNF = [
('RNF-01','Usabilidade','As ações principais devem estar acessíveis em até três interações a partir da tela inicial e apresentar confirmação visual de sucesso, espera e erro.'),
('RNF-02','Responsividade','A interface deve funcionar de 360 px a monitores desktop sem perda de conteúdo ou controles.'),
('RNF-03','Acessibilidade','Deve preservar contraste mínimo WCAG AA, foco visível, rótulos compreensíveis e navegação por teclado nas ações essenciais.'),
('RNF-04','Desempenho','Respostas internas devem atingir p95 de até 1 segundo; a busca externa de quadras deve informar carregamento e encerrar ou recorrer ao cache em até 12 segundos.'),
('RNF-05','Disponibilidade','O serviço deve expor verificação de saúde e buscar disponibilidade mensal mínima de 99% durante o período de apresentação.'),
('RNF-06','Segurança','Senhas devem ser armazenadas apenas com hash forte e salt; sessões devem ser imprevisíveis, expiradas e transportadas por HTTPS em produção.'),
('RNF-07','Privacidade e LGPD','O GPS deve ser coletado mediante consentimento, usado somente para a finalidade informada e não persistido com precisão maior ou por período superior ao necessário.'),
('RNF-08','Integridade','O banco deve impor chaves primárias, estrangeiras, unicidade, domínios e validações de faixa para impedir estados inválidos.'),
('RNF-09','Resiliência','Falhas do OpenStreetMap/Overpass não devem bloquear o restante do sistema; a aplicação deve usar cache e base de contingência.'),
('RNF-10','Compatibilidade','A aplicação deve operar nas versões atuais de Chrome, Edge, Firefox e Safari, considerando que GPS exige HTTPS fora de localhost.'),
('RNF-11','Manutenibilidade','Requisitos, regras, esquema SQL e diagrama devem permanecer versionados e identificados no repositório.'),
('RNF-12','Escalabilidade','A persistência-alvo deve ser relacional e permitir múltiplas instâncias da aplicação sem depender do disco efêmero do servidor.'),
('RNF-13','Observabilidade','Erros de API e integrações devem gerar logs com data, rota e contexto técnico, sem registrar senhas ou dados pessoais desnecessários.'),
('RNF-14','Recuperação','Em produção acadêmica persistente, devem existir cópias automáticas e restauração testada do banco antes de campeonatos e apresentações.')]

RN = [
('RN-01','O nível declarado calibra o perfil inicial, mas não corresponde ao ranking semestral.'),
('RN-02','A classificação técnica usada para equilibrar equipes deve considerar desempenho/progressão e nunca a posição no ranking do semestre.'),
('RN-03','O ranking semestral serve para classificação dos campeonatos finais de cada semestre letivo.'),
('RN-04','Treinos coletivos realizados por meio de evento em quadra concedem o dobro dos pontos semestrais do valor-base da atividade.'),
('RN-05','Treinos individuais concedem XP integral e metade do valor-base em pontos semestrais, com arredondamento para o inteiro mais próximo.'),
('RN-06','O mesmo treino individual não pode ser pontuado mais de uma vez por usuário no mesmo dia.'),
('RN-07','Treino coletivo pontuado deve estar associado a uma quadra e a evidência de participação, como inscrição/check-in.'),
('RN-08','A composição das equipes deve minimizar a diferença de classificação técnica e distribuir altura e posições; gênero é usado apenas para composição inclusiva, nunca para atribuir valor técnico.'),
('RN-09','O usuário só pode entrar em uma fila de matchmaking ativa por vez.'),
('RN-10','A partida normal prioriza equilíbrio; a rápida amplia tolerâncias de nível, horário e distância para reduzir a espera.'),
('RN-11','Uma partida oficial do MVP comporta dez jogadores, divididos em dois times de cinco; a demonstração pode formar times com quatro contas, identificando explicitamente esse modo.'),
('RN-12','A busca de quadras começa em 8 km e pode ampliar automaticamente até 25 km quando houver poucos resultados.'),
('RN-13','Sem GPS autorizado, a aplicação deve usar Torre, Recife/PE, como referência informada ao usuário e permitir nova tentativa.'),
('RN-14','Resultados externos de quadras devem ser deduplicados por fonte/identificador e distância antes da exibição.'),
('RN-15','Somente participantes confirmados podem compor uma partida; o número de inscritos não pode exceder a capacidade do evento.'),
('RN-16','Mensagens vazias ou com mais de 2.000 caracteres não são aceitas; mensagens automatizadas devem ser identificadas como bot.'),
('RN-17','Senhas nunca podem ser gravadas em texto puro nem incluídas em respostas da API, logs, documentos ou repositório.'),
('RN-18','O fechamento do ranking deve congelar a pontuação do semestre; alterações posteriores pertencem ao período seguinte ou exigem auditoria administrativa.')]

def build():
    diagram(); doc=Document(); sec=doc.sections[0]
    sec.page_width=Inches(8.5); sec.page_height=Inches(11); sec.top_margin=sec.bottom_margin=Inches(1); sec.left_margin=sec.right_margin=Inches(1); sec.header_distance=sec.footer_distance=Inches(.492)
    styles=doc.styles; normal=styles['Normal']; normal.font.name='Calibri'; normal.font.size=Pt(11); normal.font.color.rgb=RGBColor.from_string(BLACK)
    normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.1
    for name,size,before,after,color in [('Heading 1',16,16,8,ORANGE),('Heading 2',13,12,6,DARK),('Heading 3',12,8,4,DARK)]:
        s=styles[name]; s.font.name='Calibri'; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(color); s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True
    header=sec.header.paragraphs[0]; header.text='BASTREET  |  Especificação de requisitos e dados'; header.style=styles['Normal']; header.runs[0].font.size=Pt(9); header.runs[0].font.color.rgb=RGBColor.from_string(GRAY)
    add_page_number(sec.footer.paragraphs[0])
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(54); p.paragraph_format.space_after=Pt(2)
    r=p.add_run('PROJETO DE FÁBRICA DE SOFTWARE'); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=RGBColor.from_string(ORANGE)
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(4); r=p.add_run('BASTREET'); r.bold=True; r.font.size=Pt(32); r.font.color.rgb=RGBColor.from_string(BLACK)
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(22); r=p.add_run('Requisitos e Banco de Dados'); r.font.size=Pt(17); r.font.color.rgb=RGBColor.from_string(GRAY)
    add_table(doc,['INSTITUIÇÃO','DISCIPLINA'],[['UNINASSAU Olinda','Fábrica de Software']], [4680,4680])
    doc.add_heading('Integrantes', level=2)
    add_table(doc,['NOME','MATRÍCULA'],[
      ['Bárbara Menezes Nolé','01613473'],['Daniel de Moura Melo Filho','01590157'],['Ruan Deud Rameh de Oliveira','01647036']], [7000,2360])
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(14); r=p.add_run('Versão 1.0 • Recife/Olinda — 2026.2'); r.italic=True; r.font.color.rgb=RGBColor.from_string(GRAY)
    doc.add_page_break()
    doc.add_heading('1. Requisitos funcionais',level=1)
    add_table(doc,['ID','REQUISITO','DESCRIÇÃO'],RF,[900,2150,6310])
    doc.add_heading('2. Requisitos não funcionais',level=1)
    add_table(doc,['ID','ATRIBUTO','CRITÉRIO'],RNF,[900,1900,6560])
    doc.add_heading('3. Regras de negócio',level=1)
    add_table(doc,['ID','REGRA'],RN,[900,8460])
    doc.add_heading('4. Banco de dados',level=1)
    doc.add_paragraph('O modelo de referência foi definido para PostgreSQL 16 ou superior e mantém separadas a classificação técnica do matchmaking e a pontuação semestral. O esquema é uma evolução compatível e versionada: não é executado automaticamente e, portanto, não altera o funcionamento atual da aplicação.')
    doc.add_heading('4.1 Artefatos e acessos',level=2)
    links=[
      ('Esquema SQL do banco','https://github.com/Joao-NB/Bastreet/blob/main/docs/database/schema.sql'),
      ('Diagrama ER no repositório','https://github.com/Joao-NB/Bastreet/blob/main/docs/database/diagrama-er.md'),
      ('Repositório do projeto','https://github.com/Joao-NB/Bastreet'),
      ('Aplicação publicada','https://bastreet.onrender.com')]
    for label,url in links:
        p=doc.add_paragraph(style=None); p.paragraph_format.left_indent=Inches(.25); p.add_run(f'{label}: ').bold=True; hyperlink(p,url,url)
    doc.add_heading('4.2 Estrutura de dados',level=2)
    entities=[
      ('Identidade e perfil','users, player_profiles, player_availability','Autenticação, atributos esportivos, localização e horários.'),
      ('Quadras e treino','courts, training_catalog, training_sessions, court_checkins','Catálogo geográfico, conclusão de atividades e validação presencial.'),
      ('Eventos e partidas','events, event_participants, matchmaking_queue, matches, match_players','Agenda, inscrição, filas, equipes e resultado.'),
      ('Ranking','academic_semesters, ranking_entries','Pontuação isolada por semestre e fechamento do período.'),
      ('Social','conversations, conversation_members, messages','Conversas, participantes, autoria humana e bots identificados.')]
    add_table(doc,['DOMÍNIO','TABELAS','RESPONSABILIDADE'],entities,[1800,3500,4060])
    doc.add_heading('4.3 Diagrama entidade-relacionamento',level=2)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run().add_picture(str(DIAGRAM),width=Inches(6.45))
    p=doc.add_paragraph('Figura 1 — Visão agrupada das entidades do banco BASTREET. A cardinalidade completa está disponível no diagrama Mermaid do repositório.'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.runs[0].italic=True; p.runs[0].font.size=Pt(9); p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(4)
    doc.add_heading('4.4 Integridade e desempenho',level=2)
    controls=[
      ('Identificadores','UUID gerado no banco para evitar colisões entre instâncias.'),
      ('Relacionamentos','Chaves estrangeiras com exclusão em cascata ou preservação controlada conforme o domínio.'),
      ('Validações','Restrições para e-mail, altura, coordenadas, horários, pontuação, capacidade e domínio de status.'),
      ('Unicidade','E-mail, código de treino, participação, entrada no ranking e fila ativa por usuário.'),
      ('Índices','Localização, agenda de quadras, histórico de treinos, fila, ranking e mensagens cronológicas.'),
      ('Credenciais','Somente hash de senha; o esquema não contém coluna para senha em texto puro.')]
    add_table(doc,['CONTROLE','APLICAÇÃO'],controls,[2100,7260])
    doc.core_properties.title='BASTREET — Requisitos e Banco de Dados'; doc.core_properties.subject='Fábrica de Software — UNINASSAU Olinda'; doc.core_properties.author='Bárbara Menezes Nolé; Daniel de Moura Melo Filho; Ruan Deud Rameh de Oliveira'
    doc.save(OUT); print(OUT)

if __name__=='__main__': build()
